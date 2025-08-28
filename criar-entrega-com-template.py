from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def criar_documento_entrega():
    # Carrega o template
    doc = Document('template.docx')
    
    # Limpa o conteúdo existente (mantém os estilos)
    for paragraph in doc.paragraphs:
        paragraph.text = ""
    
    # Adiciona o conteúdo novo usando os estilos do template
    
    # Título Principal
    p = doc.add_paragraph('Entrega Sprint 4 - Projeto GlobalCoffee')
    p.style = 'Title'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    p = doc.add_paragraph('Consolidação de Entregas e Resultados')
    p.style = 'Subtitle'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Espaço
    
    # TASK 1
    doc.add_heading('TASK 1: CRIAÇÃO DE PERSONAS', level=1)
    doc.add_heading('Compilação de Dados e Desenho de Personas Iniciais', level=2)
    
    p = doc.add_paragraph('Foram desenvolvidas 6 personas principais para o projeto GlobalCoffee, divididas em dois grupos: Produtores de Café e Cooperativas Cafeeiras.')
    p.style = 'Normal'
    
    doc.add_paragraph()
    
    # Personas de Produtores
    p = doc.add_paragraph()
    p.add_run('Personas de Produtores:').bold = True
    
    # José Silva
    doc.add_paragraph('José Silva - Pequeno Produtor Familiar Tradicional', style='List Bullet')
    doc.add_paragraph('Perfil tecnológico básico, baixa familiaridade digital', style='List Bullet 2')
    doc.add_paragraph('Suporte intensivo no onboarding', style='List Bullet 2')
    doc.add_paragraph('Comunicação via cooperativa', style='List Bullet 2')
    
    # Roberto Santos
    doc.add_paragraph('Roberto Santos - Médio Produtor Conectado', style='List Bullet')
    doc.add_paragraph('Perfil tecnológico intermediário', style='List Bullet 2')
    doc.add_paragraph('Dashboard personalizado e ferramentas de gestão', style='List Bullet 2')
    doc.add_paragraph('Acesso a mercados premium', style='List Bullet 2')
    
    # Ana Carolina Lima
    doc.add_paragraph('Ana Carolina Lima - Jovem Produtora Inovadora', style='List Bullet')
    doc.add_paragraph('Perfil digital nativo, foco em cafés especiais', style='List Bullet 2')
    doc.add_paragraph('Ferramentas para diferenciação e construção de marca', style='List Bullet 2')
    doc.add_paragraph('Acesso a crédito jovem e modernização', style='List Bullet 2')
    
    doc.add_paragraph()
    
    # Personas de Cooperativas
    p = doc.add_paragraph()
    p.add_run('Personas de Cooperativas:').bold = True
    
    # Persona A
    doc.add_paragraph('Persona A - Pequena Cooperativa Regional', style='List Bullet')
    doc.add_paragraph('Até 1.500 membros, recepção <80.000 sacas/ano', style='List Bullet 2')
    doc.add_paragraph('Módulo financeiro plug-and-play', style='List Bullet 2')
    doc.add_paragraph('Rastreabilidade QR-Code + blockchain', style='List Bullet 2')
    
    # Persona B
    doc.add_paragraph('Persona B - Grande Cooperativa Exportadora', style='List Bullet')
    doc.add_paragraph('Exemplo: Cooxupé, >19.000 membros', style='List Bullet 2')
    doc.add_paragraph('TI corporativa avançada (Dynamics 365)', style='List Bullet 2')
    doc.add_paragraph('Nó privado Hyperledger para rastreamento', style='List Bullet 2')
    
    # Persona C
    doc.add_paragraph('Persona C - Consórcio Digital de Cooperativas', style='List Bullet')
    doc.add_paragraph('Exemplo: Supercampo, 12 cooperativas', style='List Bullet 2')
    doc.add_paragraph('Plataforma VTEX multitenant', style='List Bullet 2')
    doc.add_paragraph('Tokenização de recebíveis (Coffee Coin-like)', style='List Bullet 2')
    
    doc.add_paragraph()
    
    # TASK 2
    doc.add_heading('TASK 2: DEFINIÇÃO DO PROBLEMA', level=1)
    doc.add_heading('Insights de Dados e Pontos de Dor Identificados', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Principais Problemas Identificados:').bold = True
    
    # Volatilidade
    doc.add_paragraph('Volatilidade Extrema de Preços', style='List Bullet')
    doc.add_paragraph('Variação de R$ 250 para R$ 2.500 em 12 meses', style='List Bullet 2')
    doc.add_paragraph('Especulação financeira amplifica volatilidade em 40-60%', style='List Bullet 2')
    
    # Fragmentação
    doc.add_paragraph('Fragmentação da Cadeia Produtiva', style='List Bullet')
    doc.add_paragraph('Diversos perfis tecnológicos e regionais', style='List Bullet 2')
    doc.add_paragraph('Apenas 25% das cooperativas digitalizadas', style='List Bullet 2')
    
    # Riscos Climáticos
    doc.add_paragraph('Riscos Climáticos Crescentes', style='List Bullet')
    doc.add_paragraph('Geadas no Brasil: perdas de 30-40% da produção', style='List Bullet 2')
    doc.add_paragraph('Eventos extremos intensificados em 70% nas últimas décadas', style='List Bullet 2')
    
    # Barreiras
    doc.add_paragraph('Barreiras de Acesso a Mercados Premium', style='List Bullet')
    doc.add_paragraph('Falta de certificações ESG', style='List Bullet 2')
    doc.add_paragraph('Ausência de rastreabilidade digital', style='List Bullet 2')
    doc.add_paragraph('Certificações agregam prêmio de 15-30% no preço', style='List Bullet 2')
    
    doc.add_paragraph()
    
    # Jornada
    p = doc.add_paragraph()
    p.add_run('Jornada do Produtor Identificada:').bold = True
    
    doc.add_paragraph('1. Pré-Plantio: Análise de solo, clima e mercado', style='Normal')
    doc.add_paragraph('2. Plantio e Cultivo: Monitoramento climático contínuo', style='Normal')
    doc.add_paragraph('3. Colheita: Otimização do timing baseado em dados', style='Normal')
    doc.add_paragraph('4. Beneficiamento: Rastreabilidade e certificações', style='Normal')
    doc.add_paragraph('5. Comercialização: Acesso a múltiplos canais de venda', style='Normal')
    doc.add_paragraph('6. Pós-Venda: Análise de resultados e planejamento', style='Normal')
    
    doc.add_paragraph()
    
    # TASK 3
    doc.add_heading('TASK 3: WORKSHOP DE IDEAÇÃO', level=1)
    doc.add_heading('Materiais e Outputs do Workshop', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Metodologia Utilizada:').bold = True
    
    p = doc.add_paragraph('Foi aplicada a metodologia Duplo Diamante (Design Thinking) com as seguintes fases:')
    
    doc.add_paragraph('Descobrir: Pesquisa com 50+ produtores e 12 cooperativas', style='List Bullet')
    doc.add_paragraph('Definir: Mapeamento de 24 pontos de dor principais', style='List Bullet')
    doc.add_paragraph('Desenvolver: 120+ ideias geradas em sessões de brainstorming', style='List Bullet')
    doc.add_paragraph('Entregar: 15 funcionalidades priorizadas para MVP', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Principais Outputs do Workshop:').bold = True
    
    doc.add_paragraph('Sistema de Alertas Climáticos Simplificados', style='List Bullet')
    doc.add_paragraph('Dashboard KPIs para Cooperativas', style='List Bullet')
    doc.add_paragraph('Marketplace B2B para Café Verde', style='List Bullet')
    doc.add_paragraph('Rastreabilidade Blockchain + QR Code', style='List Bullet')
    doc.add_paragraph('API de Integração com Sistemas Legados', style='List Bullet')
    doc.add_paragraph('Mobile App Offline-First', style='List Bullet')
    
    doc.add_paragraph()
    
    # TASK 4
    doc.add_heading('TASK 4: PRIORIZAÇÃO DE FUNCIONALIDADE (MVP)', level=1)
    doc.add_heading('Problema Central e Backlog do MVP', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Problema Central Definido:').bold = True
    
    p = doc.add_paragraph('"Como conectar de forma eficiente e sustentável produtores de café de diferentes perfis tecnológicos com o mercado global, mitigando riscos climáticos e de preço, enquanto promove rastreabilidade e certificações ESG?"')
    p.style = 'Intense Quote'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Funcionalidades Priorizadas (MoSCoW):').bold = True
    
    # Tabela de priorização
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Cabeçalho
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Funcionalidade'
    hdr_cells[1].text = 'Prioridade'
    hdr_cells[2].text = 'ROI Estimado'
    hdr_cells[3].text = 'Sprint'
    
    # Aplicar negrito no cabeçalho
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Dados
    data = [
        ('API de Preços Real-time', 'Must Have', '45%', '1-2'),
        ('EUDR Compliance Scanner', 'Must Have', '60%', '7-8'),
        ('Sistema de Assinatura Digital', 'Must Have', '50%', '3-4'),
        ('Blockchain QR Code', 'Should Have', '25%', '3-4'),
        ('Integração PIX/Mobile Money', 'Should Have', '35%', '5-6'),
        ('App Mobile Offline-First', 'Must Have', '40%', '9-10'),
    ]
    
    for func, prio, roi, sprint in data:
        row_cells = table.add_row().cells
        row_cells[0].text = func
        row_cells[1].text = prio
        row_cells[2].text = roi
        row_cells[3].text = sprint
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Backlog Inicial do MVP:').bold = True
    
    doc.add_paragraph('Sistema de Autenticação e Onboarding Diferenciado', style='List Number')
    doc.add_paragraph('API de Preços com Dados de 5 Bolsas Globais', style='List Number')
    doc.add_paragraph('Alertas Climáticos com IA Preditiva', style='List Number')
    doc.add_paragraph('Dashboard KPIs para Cooperativas', style='List Number')
    doc.add_paragraph('Módulo de Certificação Digital ESG', style='List Number')
    doc.add_paragraph('Sistema de Pagamentos Integrado', style='List Number')
    doc.add_paragraph('Rastreabilidade End-to-End', style='List Number')
    doc.add_paragraph('Analytics e Relatórios Customizados', style='List Number')
    
    doc.add_paragraph()
    
    # CONCLUSÃO
    doc.add_heading('CONCLUSÃO', level=1)
    
    p = doc.add_paragraph('A Sprint 4 do projeto GlobalCoffee entregou com sucesso todos os artefatos planejados:')
    
    doc.add_paragraph('✓ 6 personas detalhadas com funcionalidades específicas', style='List Bullet')
    doc.add_paragraph('✓ Análise aprofundada dos problemas do mercado cafeeiro', style='List Bullet')
    doc.add_paragraph('✓ Workshop de ideação com 120+ ideias e 15 funcionalidades priorizadas', style='List Bullet')
    doc.add_paragraph('✓ MVP definido com backlog priorizado e roadmap de 12 meses', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph('O projeto está pronto para iniciar a fase de desenvolvimento, com foco inicial na API de preços real-time e sistema de alertas climáticos, atendendo às necessidades mais críticas identificadas durante a pesquisa.')
    
    # Salva o documento
    doc.save('entrega-sprint-4-globalcoffee-com-template.docx')
    print('Documento criado com sucesso usando o template!')

if __name__ == '__main__':
    criar_documento_entrega()