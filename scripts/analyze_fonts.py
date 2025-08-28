from docx import Document
from docx.oxml.ns import qn


def analyze_fonts(doc_path):
    doc = Document(doc_path)
    fonts = set()

    def add_font(font):
        if font and font.name:
            fonts.add(font.name)

    def collect_fonts_from_run(run):
        # Fonte direta no run
        add_font(run.font)
        # Fonte no estilo do run
        if run.style and run.style.font:
            add_font(run.style.font)

    def collect_fonts_from_paragraph(paragraph):
        # Fonte no estilo do parágrafo
        if paragraph.style and paragraph.style.font:
            add_font(paragraph.style.font)
        # Fontes nos runs do parágrafo
        for run in paragraph.runs:
            collect_fonts_from_run(run)

    # Coleta fontes dos parágrafos fora das tabelas
    for paragraph in doc.paragraphs:
        collect_fonts_from_paragraph(paragraph)

    # Coleta fontes dentro das tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    collect_fonts_from_paragraph(paragraph)

    return fonts


if __name__ == "__main__":
    fonts = analyze_fonts("Personas/Estruturas_de_Personas_updated.docx")
    print("Fonts used in document:", fonts)

quantos arquivos de python eu tenho neste projeto?
faça uma lista desses arquivos em um .md
crie um doker pra isso
